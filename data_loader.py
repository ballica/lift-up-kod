import pandas as pd
from docx import Document
import os
import logging
from config import Config
from langchain_text_splitters import RecursiveCharacterTextSplitter
logger = logging.getLogger(__name__)

class DataLoader:
    def __init__(self):
        self.data_dir = Config.DATA_DIR
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        
    def load_excel_data(self):
        """Excel dosyalarından yapılandırılmış metin verisi çıkarır."""
        documents = []
        if not os.path.exists(self.data_dir):
            logger.warning(f"Veri klasörü bulunamadı: {self.data_dir}")
            return documents

        for filename in os.listdir(self.data_dir):
            if filename.endswith(('.xlsx', '.xls')) and not filename.startswith('~$'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    # Excel'i oku, NaN değerleri boş string yap
                    df = pd.read_excel(file_path).fillna("")
                    
                    # Her satırı anlamlı bir metne dönüştür
                    text_content = f"--- DOSYA: {filename} ---\n"
                    columns = df.columns.tolist()
                    
                    for index, row in df.iterrows():
                        row_text = []
                        for col in columns:
                            if row[col]: # Boş değilse ekle
                                row_text.append(f"{col}: {row[col]}")
                        
                        if row_text:
                            text_content += " | ".join(row_text) + "\n"
                    
                    documents.append({"source": filename, "content": text_content})
                    logger.info(f"✅ Excel yüklendi: {filename}")
                except Exception as e:
                    logger.error(f"❌ Excel okuma hatası ({filename}): {str(e)}")
        return documents

    def load_word_data(self):
        """Word dosyalarından ham metin verisi çıkarır."""
        documents = []
        if not os.path.exists(self.data_dir):
            return documents

        for filename in os.listdir(self.data_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    doc = Document(file_path)
                    text = f"--- DOSYA: {filename} ---\n"
                    text += "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    
                    # Tabloları da oku
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                text += " | ".join(row_text) + "\n"
                    
                    documents.append({"source": filename, "content": text})
                    logger.info(f"✅ Word yüklendi: {filename}")
                except Exception as e:
                    logger.error(f"❌ Word okuma hatası ({filename}): {str(e)}")
        return documents

    def get_chunked_documents(self):
        """Tüm verileri yükler ve RAG için parçalar."""
        raw_docs = self.load_excel_data() + self.load_word_data()
        chunked_docs = []
        
        for doc in raw_docs:
            chunks = self.text_splitter.split_text(doc['content'])
            for chunk in chunks:
                chunked_docs.append({
                    "page_content": chunk,
                    "metadata": {"source": doc['source']}
                })
                
        logger.info(f"Toplam {len(chunked_docs)} chunk oluşturuldu.")
        return chunked_docs

    def get_dropdown_options(self):
        """Excel'den benzersiz Çalışan isimlerini ve Hedef Türlerini çeker."""
        employees = set()
        target_types = set()
        
        if not os.path.exists(self.data_dir):
            return [], []

        for filename in os.listdir(self.data_dir):
            if filename.endswith(('.xlsx', '.xls')) and not filename.startswith('~$'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    df = pd.read_excel(file_path).fillna("")
                    
                    # Kolon adı düzeltmesi: 'İsim' veya 'Ad Soyad' olabilir
                    name_col = 'İsim' if 'İsim' in df.columns else 'Ad Soyad'
                    
                    if name_col in df.columns:
                        employees.update(df[name_col].dropna().astype(str).unique())
                    
                    if 'Hedef Türü' in df.columns:
                        target_types.update(df['Hedef Türü'].dropna().astype(str).unique())
                        
                except Exception as e:
                    logger.error(f"Metadata okuma hatası ({filename}): {str(e)}")
        
        return sorted(list(employees)), sorted(list(target_types))

    def get_employee_history(self, employee_name, target_type=None):
        """Seçilen çalışan ve hedef türü için geçmiş verileri tablo olarak döner."""
        history_df = pd.DataFrame()
        
        if not os.path.exists(self.data_dir):
            return history_df

        for filename in os.listdir(self.data_dir):
            if filename.endswith(('.xlsx', '.xls')) and not filename.startswith('~$'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    df = pd.read_excel(file_path).fillna("")
                    name_col = 'İsim' if 'İsim' in df.columns else 'Ad Soyad'
                    
                    if name_col in df.columns:
                        # Çalışana göre filtrele
                        filtered = df[df[name_col] == employee_name]
                        
                        # Hedef türüne göre filtrele (opsiyonel)
                        if target_type and 'Hedef Türü' in df.columns:
                            filtered = filtered[filtered['Hedef Türü'] == target_type]
                            
                        if not filtered.empty:
                            history_df = pd.concat([history_df, filtered])
                            
                except Exception as e:
                    logger.error(f"Geçmiş verisi okuma hatası ({filename}): {str(e)}")
                    
        return history_df

    def get_employee_metadata(self, employee_name):
        """Çalışanın kimlik bilgilerini (Unvan, Bölüm, Sicil) döner."""
        metadata = {}
        
        if not os.path.exists(self.data_dir):
            return metadata

        for filename in os.listdir(self.data_dir):
            if filename.endswith(('.xlsx', '.xls')) and not filename.startswith('~$'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    df = pd.read_excel(file_path).fillna("")
                    name_col = 'İsim' if 'İsim' in df.columns else 'Ad Soyad'
                    
                    if name_col in df.columns:
                        # Çalışana göre filtrele
                        person_row = df[df[name_col] == employee_name]
                        
                        if not person_row.empty:
                            row = person_row.iloc[0]
                            # İstenen sütunlar
                            target_cols = ['Sicil', 'Unvan', 'Bölüm Ana Sorumluluk Alanı']
                            for col in target_cols:
                                if col in df.columns:
                                    metadata[col] = row[col]
                            
                            # Bulduysak çıkalım (ilk eşleşme yeterli varsayımı)
                            if metadata:
                                return metadata
                            
                except Exception as e:
                    logger.error(f"Metadata okuma hatası ({filename}): {str(e)}")
                    
        return metadata