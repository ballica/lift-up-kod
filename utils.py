from docx import Document
from docx.shared import Pt
from io import BytesIO

def generate_docx(content, title="Analiz Raporu"):
    """
    Verilen metin içeriğini Word (.docx) formatına dönüştürür.
    """
    doc = Document()
    
    # Başlık Ekle
    header = doc.add_heading(title, 0)
    
    # İçerik Ekle
    paragraph = doc.add_paragraph(content)
    run = paragraph.runs[0]
    run.font.size = Pt(11)
    
    # Bellekte PDF/Word verisini tut
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
